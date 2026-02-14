"""
TranscribeAI - 100% Local Transcription with Speaker Diarization
================================================================
Engine  : faster-whisper (CPU) or mlx-whisper (Apple Silicon GPU)
Speaker : MFCC + Agglomerative Clustering
VAD     : Silero VAD (faster-whisper) / built-in (mlx-whisper)
Language: Indonesian, English, Auto-detect (99 languages)
Input   : MP3, MP4, WAV, M4A, OGG, FLAC, WEBM
Output  : SRT, TXT, DOCX
Server  : Flask (localhost:8080)

No API key needed. Runs 100% offline.
"""

import os
import uuid
import json
import threading
import time
import sys
from datetime import datetime
from pathlib import Path

from flask import Flask, render_template, request, jsonify, send_file
import numpy as np

# ============================================================
# Engine Detection
# ============================================================
import platform

MLX_AVAILABLE = False
if platform.system() == 'Darwin' and platform.machine() == 'arm64':
    try:
        import mlx_whisper
        MLX_AVAILABLE = True
    except ImportError:
        MLX_AVAILABLE = False

FASTER_AVAILABLE = False
try:
    from faster_whisper import WhisperModel as _FWCheck
    FASTER_AVAILABLE = True
except ImportError:
    FASTER_AVAILABLE = False

# Default engine: prefer mlx on Apple Silicon
DEFAULT_ENGINE = 'mlx' if MLX_AVAILABLE else 'faster-whisper'

# ============================================================
# Global: Model cache (load once, reuse)
# ============================================================
whisper_model = None
model_lock = threading.Lock()
current_model_size = None
current_engine = None

# ============================================================
# Flask App
# ============================================================
app = Flask(__name__)

BASE_DIR = Path(__file__).parent
UPLOAD_FOLDER = BASE_DIR / 'uploads'
OUTPUT_FOLDER = BASE_DIR / 'outputs'
UPLOAD_FOLDER.mkdir(exist_ok=True)
OUTPUT_FOLDER.mkdir(exist_ok=True)

app.config['MAX_CONTENT_LENGTH'] = 2 * 1024 * 1024 * 1024  # 2GB max

ALLOWED_EXTENSIONS = {'mp3', 'mp4', 'wav', 'webm', 'm4a', 'ogg', 'flac', 'wma', 'aac'}

SPEAKER_COLORS = [
    '#4F46E5', '#DC2626', '#059669', '#D97706', '#7C3AED',
    '#DB2777', '#0891B2', '#65A30D', '#EA580C', '#6366F1'
]

MODEL_INFO = {
    'tiny':     {'params': '39M',   'vram': '~1 GB',  'speed': 'Tercepat',             'download_mb': 75},
    'base':     {'params': '74M',   'vram': '~1 GB',  'speed': 'Sangat Cepat',         'download_mb': 145},
    'small':    {'params': '244M',  'vram': '~2 GB',  'speed': 'Cepat',                'download_mb': 465},
    'medium':   {'params': '769M',  'vram': '~5 GB',  'speed': 'Sedang',               'download_mb': 1460},
    'large-v3': {'params': '1550M', 'vram': '~10 GB', 'speed': 'Lambat, Paling Akurat','download_mb': 2950},
}

# MLX model repo names (different from faster-whisper)
MLX_MODEL_REPOS = {
    'tiny':     'mlx-community/whisper-tiny-mlx',
    'base':     'mlx-community/whisper-base-mlx',
    'small':    'mlx-community/whisper-small-mlx',
    'medium':   'mlx-community/whisper-medium-mlx',
    'large-v3': 'mlx-community/whisper-large-v3-mlx',
}

# Job storage
jobs = {}


# ============================================================
# Job Class
# ============================================================
class TranscriptionJob:
    def __init__(self, job_id, filename, language, num_speakers, model_size,
                 enable_diarization=True, enable_vad=True, engine=None):
        self.job_id = job_id
        self.filename = filename
        self.language = language
        self.num_speakers = num_speakers
        self.model_size = model_size
        self.enable_diarization = enable_diarization
        self.enable_vad = enable_vad
        self.engine = engine or DEFAULT_ENGINE
        self.status = 'queued'
        self.progress = 0
        self.stage = 'Menunggu...'
        self.result = None
        self.error = None
        self.audio_path = None
        self.created_at = datetime.now().isoformat()

    def to_dict(self):
        return {
            'job_id': self.job_id,
            'filename': self.filename,
            'language': self.language,
            'model_size': self.model_size,
            'status': self.status,
            'progress': self.progress,
            'stage': self.stage,
            'result': self.result,
            'error': self.error,
            'created_at': self.created_at
        }


# ============================================================
# Helpers
# ============================================================
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def fmt_timestamp(seconds):
    """SRT format: HH:MM:SS,mmm"""
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def fmt_time(seconds):
    """Display format: MM:SS"""
    m = int(seconds // 60)
    s = int(seconds % 60)
    return f"{m:02d}:{s:02d}"


# ============================================================
# Model Loading
# ============================================================
def is_model_cached(model_size, engine=None):
    """Check if model is already downloaded in HuggingFace cache"""
    cache_dir = Path.home() / '.cache' / 'huggingface' / 'hub'

    # Check based on engine
    repo_names = []
    if engine == 'mlx' or engine is None:
        # MLX model cache: mlx-community/whisper-small-mlx → models--mlx-community--whisper-small-mlx
        mlx_repo = MLX_MODEL_REPOS.get(model_size, f"mlx-community/whisper-{model_size}-mlx")
        repo_names.append(f"models--{mlx_repo.replace('/', '--')}")
    if engine == 'faster-whisper' or engine is None:
        # faster-whisper cache: Systran/faster-whisper-small → models--Systran--faster-whisper-small
        repo_names.append(f"models--Systran--faster-whisper-{model_size}")

    for repo_name in repo_names:
        model_dir = cache_dir / repo_name
        if model_dir.exists():
            snapshots = model_dir / 'snapshots'
            if snapshots.exists() and any(snapshots.iterdir()):
                return True
    return False


def get_model(model_size='small', engine=None, job=None):
    """Load or reuse whisper model (thread-safe). Supports faster-whisper and mlx-whisper."""
    global whisper_model, current_model_size, current_engine

    if engine is None:
        engine = DEFAULT_ENGINE

    with model_lock:
        if whisper_model is not None and current_model_size == model_size and current_engine == engine:
            if job:
                job.stage = f'Model {model_size} siap (dari cache)'
                job.progress = 10
            return whisper_model

        cached = is_model_cached(model_size)
        dl_mb = MODEL_INFO.get(model_size, {}).get('download_mb', 0)

        if cached:
            msg = f'Memuat model {model_size} dari cache ({engine})...'
            print(f"  [Model] Loading '{model_size}' from cache ({engine})...")
        else:
            msg = f'Mengunduh model {model_size} (~{dl_mb}MB)... Pertama kali saja.'
            print(f"  [Model] Downloading '{model_size}' (~{dl_mb}MB)...")

        if job:
            job.stage = msg
            job.progress = 3

        t0 = time.time()

        if engine == 'mlx' and MLX_AVAILABLE:
            # mlx-whisper: use correct repo name with -mlx suffix
            repo = MLX_MODEL_REPOS.get(model_size, f"mlx-community/whisper-{model_size}-mlx")
            whisper_model = {'engine': 'mlx', 'repo': repo}
            print(f"  [Model] MLX engine ready: {repo}")
        else:
            from faster_whisper import WhisperModel
            whisper_model = WhisperModel(
                model_size,
                device="cpu",
                compute_type="int8",
                cpu_threads=os.cpu_count() or 4,
            )

        elapsed = time.time() - t0
        current_model_size = model_size
        current_engine = engine

        if job:
            job.stage = f'Model {model_size} siap! ({elapsed:.0f}s) [{engine}]'
            job.progress = 10

        print(f"  [Model] Ready in {elapsed:.1f}s ({engine}) {'(cached)' if cached else '(downloaded)'}")

        return whisper_model


# ============================================================
# Transcription (faster-whisper)
# ============================================================
def transcribe_audio(audio_path, language, model_size, job, enable_vad=True, engine=None):
    """Transcribe audio using faster-whisper or mlx-whisper"""
    if engine is None:
        engine = DEFAULT_ENGINE

    model = get_model(model_size, engine=engine, job=job)

    job.stage = f'Mentranskrip audio ({engine})...'
    job.progress = 15

    lang_param = language if language and language != 'auto' else None

    # ---- MLX-WHISPER ENGINE ----
    if engine == 'mlx' and MLX_AVAILABLE and isinstance(model, dict):
        import mlx_whisper
        repo = model['repo']

        decode_opts = {}
        if lang_param:
            decode_opts['language'] = lang_param

        job.stage = f'Mentranskrip dengan MLX (Apple Silicon GPU)...'

        result = mlx_whisper.transcribe(
            str(audio_path),
            path_or_hf_repo=repo,
            word_timestamps=False,
            **decode_opts
        )

        detected_lang = result.get('language', language or 'auto')
        segments_raw = result.get('segments', [])

        # Calculate duration from last segment
        duration = 0
        raw_segments = []
        for seg in segments_raw:
            text = seg.get('text', '').strip()
            end_t = seg.get('end', 0)
            if end_t > duration:
                duration = end_t
            if text:
                raw_segments.append({
                    'start': round(seg.get('start', 0), 2),
                    'end': round(end_t, 2),
                    'text': text,
                })

        job.progress = 60
        return raw_segments, detected_lang, duration

    # ---- FASTER-WHISPER ENGINE ----
    params = {
        'beam_size': 5,
        'word_timestamps': False,
        'vad_filter': enable_vad,
    }

    if enable_vad:
        params['vad_parameters'] = {
            'min_speech_duration_ms': 250,
            'min_silence_duration_ms': 600,
            'speech_pad_ms': 50,
        }

    if lang_param:
        params['language'] = lang_param

    segments_gen, info = model.transcribe(str(audio_path), **params)

    detected_lang = info.language
    duration = info.duration

    raw_segments = []

    for seg in segments_gen:
        text = seg.text.strip()
        if text:
            raw_segments.append({
                'start': round(seg.start, 2),
                'end': round(seg.end, 2),
                'text': text,
            })
        if duration > 0:
            pct = min(60, 15 + int(45 * seg.end / duration))
            job.progress = pct
            job.stage = f'Mentranskrip... {fmt_time(seg.end)} / {fmt_time(duration)}'

    job.progress = 60
    return raw_segments, detected_lang, duration


# ============================================================
# Speaker Diarization (MFCC + Clustering)
# ============================================================
def perform_diarization(audio_path, segments, num_speakers, job):
    """Speaker diarization using MFCC + spectral features + Agglomerative Clustering"""
    import librosa
    from sklearn.cluster import AgglomerativeClustering
    from sklearn.preprocessing import StandardScaler

    job.stage = 'Mengidentifikasi pembicara...'
    job.progress = 65

    if not segments or len(segments) < 2:
        for seg in segments:
            seg['speaker'] = 'Speaker 1'
            seg['speaker_id'] = 0
        return segments

    # Load audio (16kHz mono)
    y, sr = librosa.load(str(audio_path), sr=16000, mono=True)

    features = []
    valid_indices = []

    for i, seg in enumerate(segments):
        s0 = int(seg['start'] * sr)
        s1 = min(int(seg['end'] * sr), len(y))

        if s1 <= s0 or s0 >= len(y):
            continue

        chunk = y[s0:s1]
        if len(chunk) < int(sr * 0.3):
            continue

        try:
            # Limit analysis window to 3 seconds for performance
            analysis_chunk = chunk[:sr * 3] if len(chunk) > sr * 3 else chunk

            # MFCC (20 coefficients)
            mfcc = librosa.feature.mfcc(y=analysis_chunk, sr=sr, n_mfcc=20)
            delta = librosa.feature.delta(mfcc)
            delta2 = librosa.feature.delta(mfcc, order=2)

            # Spectral features
            sc = librosa.feature.spectral_centroid(y=analysis_chunk, sr=sr)
            sb = librosa.feature.spectral_bandwidth(y=analysis_chunk, sr=sr)
            ro = librosa.feature.spectral_rolloff(y=analysis_chunk, sr=sr)
            zcr = librosa.feature.zero_crossing_rate(analysis_chunk)

            # Pitch (YIN)
            f0 = librosa.yin(analysis_chunk, fmin=50, fmax=500, sr=sr)
            f0c = f0[f0 > 0]
            f0_mean = float(np.mean(f0c)) if len(f0c) > 0 else 0.0
            f0_std = float(np.std(f0c)) if len(f0c) > 0 else 0.0

            combined = np.vstack([mfcc, delta, delta2, sc, sb, ro, zcr])
            vec = np.concatenate([
                np.mean(combined, axis=1),
                np.std(combined, axis=1),
                [f0_mean, f0_std]
            ])

            features.append(vec)
            valid_indices.append(i)
        except Exception:
            continue

        # Update progress
        if len(segments) > 0:
            job.progress = 65 + int(15 * (i + 1) / len(segments))

    if len(features) < 2:
        for seg in segments:
            seg['speaker'] = 'Speaker 1'
            seg['speaker_id'] = 0
        return segments

    X = np.array(features)
    X_scaled = StandardScaler().fit_transform(X)

    # Auto-detect number of speakers
    if num_speakers <= 0:
        from sklearn.metrics import silhouette_score
        best_score, best_n = -1, 2
        max_n = min(6, len(X_scaled) - 1)

        for n in range(2, max_n + 1):
            try:
                lbls = AgglomerativeClustering(
                    n_clusters=n, metric='cosine', linkage='average'
                ).fit_predict(X_scaled)
                score = silhouette_score(X_scaled, lbls, metric='cosine')
                if score > best_score:
                    best_score, best_n = score, n
            except Exception:
                pass

        num_speakers = best_n
    else:
        num_speakers = min(num_speakers, len(X_scaled))

    # Cluster
    if num_speakers >= 2 and len(X_scaled) >= num_speakers:
        labels = AgglomerativeClustering(
            n_clusters=num_speakers, metric='cosine', linkage='average'
        ).fit_predict(X_scaled)
    else:
        labels = np.zeros(len(X_scaled), dtype=int)

    # Map labels to speaker numbers
    label_map = {}
    for lbl in labels:
        if lbl not in label_map:
            label_map[lbl] = len(label_map) + 1

    assigns = {}
    for idx, seg_idx in enumerate(valid_indices):
        assigns[seg_idx] = label_map[labels[idx]]

    # Assign to all segments
    for i, seg in enumerate(segments):
        if i in assigns:
            seg['speaker'] = f'Speaker {assigns[i]}'
            seg['speaker_id'] = assigns[i] - 1
        else:
            nearest = min(valid_indices, key=lambda x: abs(x - i)) if valid_indices else 0
            seg['speaker'] = f'Speaker {assigns.get(nearest, 1)}'
            seg['speaker_id'] = assigns.get(nearest, 1) - 1

    return segments


def merge_consecutive(segments):
    """Merge consecutive segments from same speaker"""
    if not segments:
        return segments
    merged = [segments[0].copy()]
    for seg in segments[1:]:
        if seg.get('speaker') == merged[-1].get('speaker'):
            merged[-1]['end'] = seg['end']
            merged[-1]['text'] += ' ' + seg['text']
        else:
            merged.append(seg.copy())
    return merged


# ============================================================
# Export Functions
# ============================================================
def generate_srt(segments, path):
    with open(path, 'w', encoding='utf-8') as f:
        for i, seg in enumerate(segments, 1):
            f.write(f"{i}\n")
            f.write(f"{fmt_timestamp(seg['start'])} --> {fmt_timestamp(seg['end'])}\n")
            sp = seg.get('speaker', '')
            f.write(f"[{sp}] {seg['text']}\n\n" if sp else f"{seg['text']}\n\n")


def generate_txt(segments, path, filename='', language='', duration=0):
    lang_names = {'id': 'Indonesian', 'en': 'English', 'auto': 'Auto-detect'}
    with open(path, 'w', encoding='utf-8') as f:
        f.write("TRANSCRIPT\n" + "=" * 60 + "\n")
        if filename:
            f.write(f"File: {filename}\n")
        f.write(f"Language: {lang_names.get(language, language)}\n")
        f.write(f"Duration: {fmt_time(duration)}\n")
        f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        speakers = sorted(set(s.get('speaker', '') for s in segments))
        f.write(f"Speakers: {', '.join(speakers)}\n")
        f.write("=" * 60 + "\n\n")

        cur_speaker = None
        for seg in segments:
            sp = seg.get('speaker', '')
            if sp != cur_speaker:
                cur_speaker = sp
                f.write(f"\n[{fmt_time(seg['start'])}] {sp}:\n")
            f.write(f"{seg['text']}\n")


def generate_docx(segments, path, filename='', language='', duration=0):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lang_names = {'id': 'Indonesian', 'en': 'English', 'auto': 'Auto-detect'}
    colors = {
        0: RGBColor(79, 70, 229), 1: RGBColor(220, 38, 38),
        2: RGBColor(5, 150, 105), 3: RGBColor(217, 119, 6),
        4: RGBColor(124, 58, 237), 5: RGBColor(219, 39, 119),
        6: RGBColor(8, 145, 178), 7: RGBColor(101, 163, 13),
    }

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = doc.add_heading('Transcript', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(30, 30, 30)

    # Metadata
    meta = []
    if filename:
        meta.append(('File', filename))
    meta.append(('Language', lang_names.get(language, language)))
    meta.append(('Duration', fmt_time(duration)))
    meta.append(('Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
    speakers = sorted(set(s.get('speaker', 'Speaker 1') for s in segments))
    meta.append(('Speakers', ', '.join(speakers)))

    for label, val in meta:
        p = doc.add_paragraph()
        r = p.add_run(f'{label}: ')
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(100, 100, 100)
        r = p.add_run(val)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(100, 100, 100)
        p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph('_' * 70).runs[0].font.color.rgb = RGBColor(200, 200, 200)

    # Transcript
    for seg in segments:
        p = doc.add_paragraph()
        r = p.add_run(f'[{fmt_time(seg["start"])}]  ')
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(150, 150, 150)

        sp = seg.get('speaker', '')
        sid = seg.get('speaker_id', 0)
        if sp:
            r = p.add_run(f'{sp}\n')
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = colors.get(sid % len(colors), RGBColor(79, 70, 229))

        r = p.add_run(seg['text'])
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(30, 30, 30)
        r.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(14)

    doc.save(str(path))


# ============================================================
# Job Processor
# ============================================================
def process_job(job_id):
    job = jobs.get(job_id)
    if not job:
        return

    try:
        file_path = UPLOAD_FOLDER / f"{job_id}_{job.filename}"

        # 1. Model Loading
        job.status = 'processing'
        job.progress = 2
        get_model(job.model_size, job=job)

        # 2. Transcription
        job.stage = 'Mentranskrip audio...'
        job.progress = 10
        job.audio_path = str(file_path)

        segments, detected_lang, duration = transcribe_audio(
            file_path, job.language, job.model_size, job, job.enable_vad, job.engine
        )

        if not segments:
            raise Exception("Tidak ada teks yang terdeteksi dari audio.")

        # 3. Speaker Diarization
        if job.enable_diarization and len(segments) >= 2:
            segments = perform_diarization(file_path, segments, job.num_speakers, job)
            segments = merge_consecutive(segments)
        else:
            for seg in segments:
                seg['speaker'] = 'Speaker 1'
                seg['speaker_id'] = 0

        job.progress = 85

        # 4. Generate output files
        job.stage = 'Membuat file output...'
        out_base = str(OUTPUT_FOLDER / job_id)

        srt_path = f"{out_base}.srt"
        generate_srt(segments, srt_path)

        txt_path = f"{out_base}.txt"
        generate_txt(segments, txt_path, job.filename, detected_lang, duration)

        docx_path = f"{out_base}.docx"
        generate_docx(segments, docx_path, job.filename, detected_lang, duration)

        job.progress = 95

        # 5. Build result
        job.result = {
            'segments': segments,
            'duration': duration,
            'language': detected_lang,
            'model': job.model_size,
            'num_speakers': len(set(s.get('speaker', '') for s in segments)),
            'total_segments': len(segments),
            'files': {
                'srt': srt_path,
                'txt': txt_path,
                'docx': docx_path,
            },
            'speaker_colors': SPEAKER_COLORS,
        }

        job.status = 'completed'
        job.stage = 'Selesai!'
        job.progress = 100

    except Exception as e:
        job.status = 'error'
        job.error = str(e)
        job.stage = f'Error: {str(e)}'
        import traceback
        traceback.print_exc()


# ============================================================
# Routes
# ============================================================
@app.route('/')
def index():
    return render_template('index.html', models=MODEL_INFO)


@app.route('/api/upload', methods=['POST'])
def upload():
    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400

    file = request.files['file']
    if not file.filename or not allowed_file(file.filename):
        exts = ', '.join(ALLOWED_EXTENSIONS)
        return jsonify({'error': f'Format tidak didukung. Gunakan: {exts}'}), 400

    language = request.form.get('language', 'auto')
    num_speakers = int(request.form.get('num_speakers', 0))
    model_size = request.form.get('model_size', 'small')
    enable_diarization = request.form.get('enable_diarization', 'true') == 'true'
    enable_vad = request.form.get('enable_vad', 'true') == 'true'
    engine = request.form.get('engine', DEFAULT_ENGINE)

    job_id = str(uuid.uuid4())[:8]
    job = TranscriptionJob(
        job_id, file.filename, language, num_speakers,
        model_size, enable_diarization, enable_vad, engine
    )
    jobs[job_id] = job

    file_path = UPLOAD_FOLDER / f"{job_id}_{file.filename}"
    file.save(str(file_path))

    thread = threading.Thread(target=process_job, args=(job_id,))
    thread.daemon = True
    thread.start()

    return jsonify({'job_id': job_id, 'filename': file.filename})


@app.route('/api/status/<job_id>')
def status(job_id):
    job = jobs.get(job_id)
    if not job:
        return jsonify({'error': 'Not found'}), 404
    return jsonify(job.to_dict())


@app.route('/api/download/<job_id>/<fmt>')
def download(job_id, fmt):
    job = jobs.get(job_id)
    if not job or job.status != 'completed':
        return jsonify({'error': 'Not ready'}), 404

    fpath = job.result['files'].get(fmt)
    if not fpath or not os.path.exists(fpath):
        return jsonify({'error': 'File not found'}), 404

    base = job.filename.rsplit('.', 1)[0]
    mimes = {
        'srt': 'application/x-subrip',
        'txt': 'text/plain',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    }
    return send_file(fpath, as_attachment=True,
                     download_name=f"{base}_transcript.{fmt}",
                     mimetype=mimes.get(fmt, 'application/octet-stream'))


@app.route('/api/audio/<job_id>')
def audio(job_id):
    job = jobs.get(job_id)
    if not job or not job.audio_path:
        return jsonify({'error': 'No audio'}), 404
    return send_file(job.audio_path)


@app.route('/api/models')
def list_models():
    return jsonify(MODEL_INFO)


@app.route('/api/engine-info')
def engine_info():
    """Return available engines and current default"""
    engines = []
    if FASTER_AVAILABLE:
        engines.append({
            'id': 'faster-whisper',
            'name': 'faster-whisper (CPU)',
            'desc': 'CTranslate2 optimized, works on all systems',
            'available': True,
        })
    if MLX_AVAILABLE:
        engines.append({
            'id': 'mlx',
            'name': 'mlx-whisper (Apple Silicon GPU)',
            'desc': 'Optimized for M1/M2/M3/M4, 2-5x lebih cepat',
            'available': True,
        })
    if not engines:
        engines.append({
            'id': 'faster-whisper',
            'name': 'faster-whisper (not installed)',
            'desc': 'Run: pip install faster-whisper',
            'available': False,
        })
    return jsonify({
        'default': DEFAULT_ENGINE,
        'engines': engines,
        'apple_silicon': platform.system() == 'Darwin' and platform.machine() == 'arm64',
    })


@app.route('/api/model-status/<model_size>')
def model_status(model_size):
    """Check if a model is already downloaded"""
    cached = is_model_cached(model_size)
    info = MODEL_INFO.get(model_size, {})
    return jsonify({
        'model': model_size,
        'cached': cached,
        'download_mb': info.get('download_mb', 0),
        'params': info.get('params', ''),
        'speed': info.get('speed', ''),
    })


# ============================================================
# Main
# ============================================================
if __name__ == '__main__':
    print()
    print("=" * 58)
    print("  TranscribeAI - 100% Local Transcription")
    print("=" * 58)
    print()
    if MLX_AVAILABLE:
        print("  Engine  : mlx-whisper (Apple Silicon GPU) ⚡")
        print("            + faster-whisper (CPU fallback)")
    elif FASTER_AVAILABLE:
        print("  Engine  : faster-whisper (CPU)")
    else:
        print("  Engine  : NONE (install faster-whisper or mlx-whisper)")
    print()
    print("  URL     : http://localhost:8080")
    print("  Models  : tiny, base, small, medium, large-v3")
    print("  Language: Indonesian, English, 99+ bahasa")
    print("  Input   : MP3, MP4, WAV, M4A, OGG, FLAC")
    print("  Output  : SRT, TXT, DOCX")
    print()
    print("  Tanpa API key! 100% offline & gratis.")
    print("  Tekan Ctrl+C untuk berhenti")
    print("=" * 58)
    print()

    # Pre-load default model
    if '--preload' in sys.argv:
        print("  Pre-loading model 'small'...")
        get_model('small')

    app.run(host='0.0.0.0', port=8080, debug=False)
