"""
TranscribeAI - CLI Version (faster-whisper)
Process audio files directly without web server.
Usage: python3 transcribe_cli.py <audio_file> [--language id|en|auto] [--speakers N] [--model small]

100% Local - No API key needed.
"""

import os
import sys
import time
from datetime import datetime
from pathlib import Path

import numpy as np


SPEAKER_COLORS_NAMES = ['Biru', 'Merah', 'Hijau', 'Kuning', 'Ungu', 'Pink', 'Cyan', 'Lime']


def log(msg):
    print(f"  [{time.strftime('%H:%M:%S')}] {msg}")


def format_timestamp(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds % 1) * 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"


def format_time(seconds):
    minutes = int(seconds // 60)
    secs = int(seconds % 60)
    return f"{minutes:02d}:{secs:02d}"


def transcribe(audio_path, language='auto', model_size='small'):
    from faster_whisper import WhisperModel

    log(f"Memuat model '{model_size}' (CPU, int8)...")
    model = WhisperModel(
        model_size,
        device="cpu",
        compute_type="int8",
        cpu_threads=os.cpu_count() or 4,
    )

    log("Mentranskrip dengan faster-whisper...")
    params = {
        'beam_size': 5,
        'word_timestamps': False,
        'vad_filter': True,
        'vad_parameters': {
            'min_speech_duration_ms': 250,
            'min_silence_duration_ms': 600,
            'speech_pad_ms': 50,
        },
    }

    if language and language != 'auto':
        params['language'] = language

    segments_gen, info = model.transcribe(str(audio_path), **params)

    detected_lang = info.language
    duration = info.duration

    segments = []
    for seg in segments_gen:
        text = seg.text.strip()
        if text:
            segments.append({
                'start': round(seg.start, 2),
                'end': round(seg.end, 2),
                'text': text,
            })

    log(f"Transkripsi selesai: {len(segments)} segmen, bahasa: {detected_lang}")
    return segments, detected_lang, duration


def diarize(audio_path, segments, num_speakers=0):
    log("Mengidentifikasi pembicara...")
    import librosa
    from sklearn.cluster import AgglomerativeClustering
    from sklearn.preprocessing import StandardScaler

    if not segments or len(segments) < 2:
        for s in segments:
            s['speaker'] = 'Speaker 1'
            s['speaker_id'] = 0
        return segments

    y, sr = librosa.load(str(audio_path), sr=16000, mono=True)
    features, valid_idx = [], []

    for i, seg in enumerate(segments):
        s0 = int(seg['start'] * sr)
        s1 = min(int(seg['end'] * sr), len(y))
        if s1 <= s0 or s0 >= len(y):
            continue
        chunk = y[s0:s1]
        if len(chunk) < int(sr * 0.3):
            continue
        try:
            analysis_chunk = chunk[:sr * 3] if len(chunk) > sr * 3 else chunk

            mfcc = librosa.feature.mfcc(y=analysis_chunk, sr=sr, n_mfcc=20)
            delta = librosa.feature.delta(mfcc)
            delta2 = librosa.feature.delta(mfcc, order=2)
            sc = librosa.feature.spectral_centroid(y=analysis_chunk, sr=sr)
            sb = librosa.feature.spectral_bandwidth(y=analysis_chunk, sr=sr)
            sr_ = librosa.feature.spectral_rolloff(y=analysis_chunk, sr=sr)
            zcr = librosa.feature.zero_crossing_rate(analysis_chunk)
            f0 = librosa.yin(analysis_chunk, fmin=50, fmax=500, sr=sr)
            f0c = f0[f0 > 0]
            combined = np.vstack([mfcc, delta, delta2, sc, sb, sr_, zcr])
            vec = np.concatenate([
                np.mean(combined, axis=1), np.std(combined, axis=1),
                [np.mean(f0c) if len(f0c) > 0 else 0, np.std(f0c) if len(f0c) > 0 else 0]
            ])
            features.append(vec)
            valid_idx.append(i)
        except Exception:
            continue

    if len(features) < 2:
        for s in segments:
            s['speaker'] = 'Speaker 1'
            s['speaker_id'] = 0
        return segments

    X = StandardScaler().fit_transform(np.array(features))

    if num_speakers <= 0:
        from sklearn.metrics import silhouette_score
        best_s, best_n = -1, 2
        for n in range(2, min(7, len(X))):
            try:
                lbl = AgglomerativeClustering(n_clusters=n, metric='cosine', linkage='average').fit_predict(X)
                sc = silhouette_score(X, lbl, metric='cosine')
                if sc > best_s:
                    best_s, best_n = sc, n
            except Exception:
                pass
        num_speakers = best_n

    num_speakers = min(num_speakers, len(X))
    if num_speakers >= 2:
        labels = AgglomerativeClustering(n_clusters=num_speakers, metric='cosine', linkage='average').fit_predict(X)
    else:
        labels = np.zeros(len(X), dtype=int)

    lmap = {}
    for l in labels:
        if l not in lmap:
            lmap[l] = len(lmap) + 1

    assigns = {valid_idx[i]: lmap[labels[i]] for i in range(len(labels))}

    for i, seg in enumerate(segments):
        if i in assigns:
            seg['speaker'] = f'Speaker {assigns[i]}'
            seg['speaker_id'] = assigns[i] - 1
        else:
            nearest = min(valid_idx, key=lambda x: abs(x - i)) if valid_idx else 0
            seg['speaker'] = f'Speaker {assigns.get(nearest, 1)}'
            seg['speaker_id'] = assigns.get(nearest, 1) - 1

    n = len(set(s['speaker'] for s in segments))
    log(f"Terdeteksi {n} pembicara")
    return segments


def merge_segments(segments):
    if not segments:
        return segments
    merged = [segments[0].copy()]
    for seg in segments[1:]:
        if seg.get('speaker') == merged[-1].get('speaker'):
            merged[-1]['end'] = seg['end']
            merged[-1]['text'] += ' ' + seg['text']
        else:
            merged.append(seg.copy())
    return merged


def save_srt(segments, path):
    with open(path, 'w', encoding='utf-8') as f:
        for i, s in enumerate(segments, 1):
            f.write(f"{i}\n{format_timestamp(s['start'])} --> {format_timestamp(s['end'])}\n")
            f.write(f"[{s.get('speaker', '')}] {s['text']}\n\n")
    log(f"SRT: {path}")


def save_txt(segments, path, filename='', language='', duration=0):
    lang_names = {'id': 'Indonesian', 'en': 'English', 'auto': 'Auto-detect'}
    with open(path, 'w', encoding='utf-8') as f:
        f.write("TRANSCRIPT\n" + "=" * 60 + "\n")
        if filename:
            f.write(f"File: {filename}\n")
        f.write(f"Language: {lang_names.get(language, language)}\n")
        f.write(f"Duration: {format_time(duration)}\n")
        f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write("=" * 60 + "\n\n")
        cur = None
        for s in segments:
            sp = s.get('speaker', '')
            if sp != cur:
                cur = sp
                f.write(f"\n[{format_time(s['start'])}] {sp}:\n")
            f.write(f"{s['text']}\n")
    log(f"TXT: {path}")


def save_docx(segments, path, filename='', language='', duration=0):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    lang_names = {'id': 'Indonesian', 'en': 'English', 'auto': 'Auto-detect'}
    colors = {
        0: RGBColor(79, 70, 229), 1: RGBColor(220, 38, 38),
        2: RGBColor(5, 150, 105), 3: RGBColor(217, 119, 6),
        4: RGBColor(124, 58, 237), 5: RGBColor(219, 39, 119),
        6: RGBColor(8, 145, 178), 7: RGBColor(101, 163, 13),
    }

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title = doc.add_heading('Transcript', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = [('File', filename), ('Language', lang_names.get(language, language)),
            ('Duration', format_time(duration)),
            ('Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ('Speakers', ', '.join(sorted(set(s.get('speaker', 'Speaker 1') for s in segments))))]
    for label, val in meta:
        if val:
            p = doc.add_paragraph()
            r = p.add_run(f'{label}: ')
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(100, 100, 100)
            r = p.add_run(val)
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(100, 100, 100)
            p.paragraph_format.space_after = Pt(2)

    doc.add_paragraph('_' * 70).runs[0].font.color.rgb = RGBColor(200, 200, 200)

    for seg in segments:
        p = doc.add_paragraph()
        r = p.add_run(f'[{format_time(seg["start"])}]  ')
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(150, 150, 150)
        sp = seg.get('speaker', '')
        sid = seg.get('speaker_id', 0)
        if sp:
            r = p.add_run(f'{sp}\n')
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = colors.get(sid % len(colors), RGBColor(79, 70, 229))
        r = p.add_run(seg['text'])
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(30, 30, 30)
        p.paragraph_format.space_after = Pt(14)

    doc.save(str(path))
    log(f"DOCX: {path}")


def main():
    import argparse
    parser = argparse.ArgumentParser(description='TranscribeAI - CLI Transcription (faster-whisper)')
    parser.add_argument('file', help='Audio/video file (MP3, MP4, WAV, etc.)')
    parser.add_argument('--language', '-l', default='auto', choices=['auto', 'id', 'en'],
                        help='Language: auto, id (Indonesian), en (English)')
    parser.add_argument('--speakers', '-s', type=int, default=0,
                        help='Number of speakers (0 = auto-detect)')
    parser.add_argument('--model', '-m', default='small',
                        choices=['tiny', 'base', 'small', 'medium', 'large-v3'],
                        help='Model size (default: small)')
    parser.add_argument('--no-diarization', action='store_true',
                        help='Disable speaker diarization')
    parser.add_argument('--output', '-o', default=None,
                        help='Output directory (default: same as input)')
    args = parser.parse_args()

    if not os.path.exists(args.file):
        print(f"ERROR: File tidak ditemukan: {args.file}")
        sys.exit(1)

    filename = os.path.basename(args.file)
    base = filename.rsplit('.', 1)[0]
    out_dir = args.output or os.path.dirname(os.path.abspath(args.file))

    print()
    print("=" * 55)
    print("  TranscribeAI - CLI Transcription")
    print("  Engine: faster-whisper (100% Local)")
    print("=" * 55)
    print(f"  File    : {filename}")
    print(f"  Model   : {args.model}")
    print(f"  Bahasa  : {args.language}")
    print(f"  Speaker : {'Auto' if args.speakers == 0 else args.speakers}")
    print("=" * 55)
    print()

    t0 = time.time()

    # Transcribe
    segments, detected_lang, duration = transcribe(args.file, args.language, args.model)

    # Diarize
    if not args.no_diarization and len(segments) >= 2:
        segments = diarize(args.file, segments, args.speakers)
        segments = merge_segments(segments)
    else:
        for s in segments:
            s['speaker'] = 'Speaker 1'
            s['speaker_id'] = 0

    # Save outputs
    log("Menyimpan file output...")
    srt_path = os.path.join(out_dir, f"{base}_transcript.srt")
    txt_path = os.path.join(out_dir, f"{base}_transcript.txt")
    docx_path = os.path.join(out_dir, f"{base}_transcript.docx")

    save_srt(segments, srt_path)
    save_txt(segments, txt_path, filename, detected_lang, duration)
    save_docx(segments, docx_path, filename, detected_lang, duration)

    elapsed = time.time() - t0

    print()
    print("=" * 55)
    print("  SELESAI!")
    print("=" * 55)
    print(f"  Waktu   : {elapsed:.1f} detik")
    print(f"  Durasi  : {format_time(duration)}")
    print(f"  Segmen  : {len(segments)}")
    print(f"  Speaker : {len(set(s['speaker'] for s in segments))}")
    print(f"  Bahasa  : {detected_lang}")
    print("-" * 55)
    print(f"  SRT  : {srt_path}")
    print(f"  TXT  : {txt_path}")
    print(f"  DOCX : {docx_path}")
    print("=" * 55)
    print()

    # Print transcript preview
    print("PREVIEW TRANSKRIP:")
    print("-" * 55)
    for seg in segments[:10]:
        t = format_time(seg['start'])
        sp = seg.get('speaker', '')
        print(f"  [{t}] {sp}: {seg['text'][:80]}{'...' if len(seg['text']) > 80 else ''}")
    if len(segments) > 10:
        print(f"  ... dan {len(segments) - 10} segmen lainnya")
    print("-" * 55)


if __name__ == '__main__':
    main()
